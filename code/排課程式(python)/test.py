from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
import pandas as pd
from docx.shared import Inches
from docx import Document

df = pd.DataFrame({"col_1": [1, 2, 3, 4, 5,], "col_2": [6, 7, 8, 9, 10]})
df["fruit"] = ["apple", "guava", "banana", "cherry", "mango"]
df["color"] = ["black", "green", "black", "green", "black"]

fruits = ["apple", "guava", "banana", "cherry", "mango"]
colors = ["black", "green"]
df["fruit_i"] = pd.Categorical(df['fruit'], categories=fruits, ordered=True)
df["color_i"] = pd.Categorical(df['color'], categories=colors, ordered=True)
df.sort_values(by=['fruit_i', 'color_i'], inplace=True)
df.set_index(keys=['fruit_i', 'color_i'], inplace=True)
df2 = df.T
# print(df2.index[0])
# print(df2.iloc[2, 2])
print(df2.columns[0][0])


def dataframe_to_word(df, filename):  # 列為星期，開課單位
    document = Document()
    for section in document.sections:  # 將頁面設為A3橫向
        section.page_width = Inches(16.5)
        section.page_height = Inches(11.7)
        section.header_distance = Inches(1.27 * 0.393701)
        section.footer_distance = Inches(1.27 * 0.393701)
        section.top_margin = Inches(0.4 * 0.393701)  # 0.4cm
        section.bottom_margin = Inches(0.1 * 0.393701)  # 0.1cm
        section.left_margin = Inches(0.1 * 0.393701)
        section.right_margin = Inches(0.1 * 0.393701)

    table = document.add_table(rows=df.shape[0] + 2, cols=df.shape[1] + 1)
    # 添加索引名稱
    table.cell(0, 0).text = df.columns.names[0]
    table.cell(1, 0).text = df.columns.names[1]
    for j in range(df.shape[0]):  # 添加列標題(rows)
        table.cell(j+2, 0).text = df.index[j]
    for i in range(df.shape[1]):  # 添加欄標題(columns) #i:0~4
        for j in range(2):  # j:0,1
            index_value = df.columns[i][j]
            table.cell(j, i+1).text = str(index_value)
    for i in range(df.shape[0]):
        for j in range(df.shape[1]):
            table.cell(i + 2, j + 1).text = str(df.iloc[i, j])  # 需在仔細檢查

    document.save(filename)


dataframe_to_word(df2, r"C:/Users/User/OneDrive/Desktop/排課程式/嗨.docx")

document = Document()
document.save(r"C:/Users/User/OneDrive/Desktop/排課程式/test_test.docx")

# note


def set_cell_width(cell, width):
    tcPr = cell._element.get_or_add_tcPr()  # 获取或添加单元格属性
    tcW = OxmlElement('w:tcW')  # 创建宽度属性元素
    tcW.set(qn('w:w'), str(width))  # 设置宽度值
    tcW.set(qn('w:type'), 'dxa')  # 设置宽度类型
    tcPr.append(tcW)  # 将宽度属性添加到单元格属性中


def set_row_height(row, height):
    trPr = row._tr.get_or_add_trPr()  # 获取或添加行属性
    trHeight = OxmlElement('w:trHeight')  # 创建高度属性元素
    trHeight.set(qn('w:val'), str(height))  # 设置高度值
    trPr.append(trHeight)  # 将高度属性添加到行属性中


def dataframe_to_word(df, filename):
    document = Document()

    # 将页面设为A3横向
    for section in document.sections:
        section.page_width = Inches(16.5)
        section.page_height = Inches(11.7)
        section.header_distance = Inches(1.27 * 0.393701)
        section.footer_distance = Inches(1.27 * 0.393701)
        section.top_margin = Inches(0.4 * 0.393701)  # 0.4cm
        section.bottom_margin = Inches(0.1 * 0.393701)  # 0.1cm
        section.left_margin = Inches(0.1 * 0.393701)
        section.right_margin = Inches(0.1 * 0.393701)

    # 创建表格，行数为df的行数+2（用于索引名称和列标题），列数为df的列数+1（用于行索引）
    table = document.add_table(rows=df.shape[0] + 2, cols=df.shape[1] + 1)

    # 设置列宽
    for col in table.columns:
        for cell in col.cells:
            # 设置每个单元格的宽度为2000 twips（1英寸 = 1440 twips）
            set_cell_width(cell, 2000)

    # 设置行高
    for row in table.rows:
        set_row_height(row, 400)  # 设置每行的高度为400 twips

    # 添加索引名称
    if df.columns.names[0] is not None:
        table.cell(0, 0).text = df.columns.names[0]
    if df.columns.names[1] is not None:
        table.cell(1, 0).text = df.columns.names[1]

    # 添加列标题
    for i, col in enumerate(df.columns):
        table.cell(0, i + 1).text = str(col[0])
        table.cell(1, i + 1).text = str(col[1])

    # 添加行标题
    for j in range(df.shape[0]):
        table.cell(j + 2, 0).text = str(df.index[j])

    # 添加数据
    for i in range(df.shape[0]):
        for j in range(df.shape[1]):
            table.cell(i + 2, j + 1).text = str(df.iloc[i, j])

    # 保存文件
    document.save(filename)


# 示例 DataFrame
data = {
    '開課單位': ['數學系', '數學系', '數學系'],
    '課程名稱': ['微積分', '線性代數', '概率論'],
    '主授教師': ['張三', '李四', '王五'],
    '週次': [1, 1, 1],
    '星期': [1, 2, 3],
    '節次': [1, 2, 3],
    '教室': ['A101', 'A102', 'A103'],
    '星期.1': [1, 2, 3],
    '節次.1': [1, 2, 3],
    '教室.1': ['A101', 'A102', 'A103']
}
df = pd.DataFrame(data)
df.set_index(['課程名稱', '主授教師'], inplace=True)
df.columns.names = ['星期', '開課單位']

# 使用示例 DataFrame 调用函数并保存 Word 文件
dataframe_to_word(df, r'C:\Users\User\OneDrive\Desktop\排課程式\數學系課表.docx')

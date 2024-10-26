# 請確認必修/必選科目清單是否需更新(required_sub)
# 請確認課程資訊是否依照正確格式填寫
# 請確認程式中的檔案路徑是否需更改
# 請確認是否安裝必要套件
# E1到E3之課程不列入課表，需另外註記

import os
from docx.shared import Inches
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
import pandas as pd


current_directory = os.path.dirname(os.path.abspath(__file__))
excel_file_path = os.path.join(current_directory, '112-2數學系課程資訊.xlsx')
course = pd.read_excel(excel_file_path, header=1, sheet_name=None)


# 必修課程清單
required_sub = ['資訊數學導論', '線性代數(一)', '線性代數(二)', '高等微積分(一)', '高等微積分(二)', '代數學(一)', '離散數學', '統計學', '科學計算',
                '演算法', '資料結構', '機率論', '微積分(二)', '人工智慧概論', '程式設計(ㄧ)', '微積分(ㄧ)', '數學導論', '數值分析', '微分方程', '複變函數論']


# course = pd.read_excel(r"C:\Users\User\OneDrive\Desktop\排課程式\112-2數學系課程資訊.xlsx", header=1, sheet_name=None)

df1 = course['D1B-數學系資訊數學組'][['開課單位', '課程名稱', '主授教師',
                              '週次', '星期', '節次', '教室', '星期.1', '節次.1', '教室.1']]
df2 = course['D31-數學系應用數學組'][['開課單位', '課程名稱', '主授教師',
                              '週次', '星期', '節次', '教室', '星期.1', '節次.1', '教室.1']]

df1 = df1.drop(index=0)
df2 = df2.drop(index=0)
df1.fillna(0, inplace=True)
df2.fillna(0, inplace=True)


def simp(df):  # 將課表之星期.1，節次.1，教室.1欄位之資料保存後，去除欄位
    temp = pd.DataFrame()
    x = pd.DataFrame()
    for i in range(len(df.index)):
        if df.iloc[i, 8] != 0:  # 星期.1有值，四學分的課的情形
            x = df.iloc[i, :]
            x.iloc[4:7] = x.iloc[7:10].values
            x = pd.DataFrame(x).transpose()
            temp = pd.concat([temp, x])
    df = pd.concat([df, temp])
    df = df.drop(columns=['星期.1', '節次.1', '教室.1'])
    df = df.reset_index(drop=True)
    return df


df1 = simp(df1)
df2 = simp(df2)


for i in [1, 2]:
    s = 'df' + str(i) + "['主授教師'] = '(' + df" + \
        str(i) + "['主授教師'].astype(str) + ')'"
    exec(s)
# dfi['主授教師'] = '(' + dfi['主授教師'].astype(str) + ')'


def aggre(df):  # 欲輸出課表之欄位值
    info = []
    for i in range(len(df.index)):
        x = str(df.iloc[i, 1])+str(df.iloc[i, 2])+'-'+str(df.iloc[i, 6])
        info.append(x)
    df["info"] = info
    return (df)


df1 = aggre(df1)
df2 = aggre(df2)

df_math = pd.concat([df1, df2])

# 以星期，開課單位作為索引
days = ['一', '二', '三', '四', '五']
grades_1 = ['資數一', '應數一', '資數二', '應數二', '資數三', '應數三', '資數四',
            '應數四', '資數組軍訓', '應數組軍訓', '資數組人哲', '應數組人哲', '資數組專倫', '應數組專倫']
# 應數人哲:應二
# 資數人哲:資二
# 應數專倫:應三
# 資數專倫:資三(暫定)
# 應數軍訓:應一
# 資數軍訓:資一


df_math['星期'] = pd.Categorical(df_math['星期'], categories=days, ordered=True)
df_math['開課單位'] = pd.Categorical(
    df_math['開課單位'], categories=grades_1, ordered=True)

df_math.sort_values(by=['星期', '開課單位'], inplace=True)
df_math.set_index(keys=['星期', '開課單位'], inplace=True)
df_math = df_math.rename(index={'應數組人哲': '應數二'}, level=1)
df_math = df_math.rename(index={'資數組人哲': '資數二'}, level=1)
df_math = df_math.rename(index={'應數組專倫': '應數三'}, level=1)
df_math = df_math.rename(index={'資數組專倫': '資數三'}, level=1)
df_math = df_math.rename(index={'應數組軍訓': '資數一'}, level=1)
df_math = df_math.rename(index={'資數組軍訓': '資數一'}, level=1)
df_math.index = df_math.index.set_levels(
    df_math.index.levels[1].str.replace('數', ''), level=1)
df_math['start_time'] = df_math['節次'].str.slice(0, 2)
df_math['end_time'] = df_math['節次'].str.slice(-2)

# 將課程時間依從早到晚的順序編號
time = ['D1', 'D2', 'D3', 'D4', 'DN', 'D5',
        'D6', 'D7', 'D8', 'E0', 'E1', 'E2', 'E3']
df_math['start_index'] = df_math['start_time'].apply(lambda x: time.index(x))
df_math['end_index'] = df_math['end_time'].apply(lambda x: time.index(x))

# 刪除'E1','E2','E3'的課
df_math = df_math[df_math['start_index'] < 10]
df_math = df_math[df_math['end_index'] < 10]


x = list(range(df_math.shape[0]))
for i in range(df_math.shape[0]):
    k = df_math.iloc[i, 8]  # i-th row start_index
    j = df_math.iloc[i, 9]  # i-th row end_index
    x[i] = list(range(k, j+1))
df_math['class_time'] = x

# 標記需先填入課表的課
df_math['priority'] = 0
for i in range(df_math.shape[0]):
    if df_math.iloc[i, 0] in required_sub:
        df_math.iloc[i, 11] = 1

df_math = df_math.drop(columns=['課程名稱', '主授教師', '週次', '節次',
                       '教室', 'start_time', 'end_time', 'start_index', 'end_index'])


df_math = df_math.T
# df_math.to_excel(r"C:\Users\User\OneDrive\Desktop\排課程式\數學系課表測試2.xlsx", index=True)
# column name也可具有多個層級，如同multiindex


# 輸入之dataframe，其column需具multiindex
def dataframe_to_word(df, filename):  # 行名為[星期，開課單位]
    document = Document()

    # 添加標題
    today = datetime.today()
    roc_year = today.year - 1911
    date = f"{roc_year}.{today.month}.{today.day}"
    title_name = '輔仁大學數學系112學年度第二學期課表' + date
    title = document.add_heading(title_name, level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.runs[0]
    run.font.size = Pt(16)
    run.font.name = '標楷體'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
    run.font.color.rgb = RGBColor(0, 0, 0)

    def set_cell_width(cell, width):  # 設置儲存格欄寬
        tcPr = cell._element.get_or_add_tcPr()  # 獲取儲存格屬性
        tcW = OxmlElement('w:tcW')  # 新增XML屬性
        tcW.set(qn('w:w'), str(width))  # 設置寬度
        tcW.set(qn('w:type'), 'dxa')  # 設置寬度類別:絕對寬度
        tcPr.append(tcW)

    def set_row_height(row, height):
        trPr = row._tr.get_or_add_trPr()  # 獲取row屬性
        trHeight = OxmlElement('w:trHeight')  # 新增XML屬性:高度
        trHeight.set(qn('w:val'), str(height))  # 設置高度
        trPr.append(trHeight)

    for section in document.sections:  # 將頁面設為A3橫向
        section.page_width = Inches(16.5)
        section.page_height = Inches(11.7)
        section.header_distance = Inches(1.27 * 0.393701)
        section.footer_distance = Inches(1.27 * 0.393701)
        section.top_margin = Inches(0.4 * 0.393701)  # 0.4cm
        section.bottom_margin = Inches(0.1 * 0.393701)  # 0.1cm
        section.left_margin = Inches(0.1 * 0.393701)
        section.right_margin = Inches(0.1 * 0.393701)

    table = document.add_table(rows=12, cols=42)  # 僅記錄到E0之課程

    for rows in table.rows:
        for cell in rows.cells[1:]:
            set_cell_width(cell, 510)  # 設置儲存格寬度為0.9cm

    row_height = [0.55, 0.89, 2.32, 3, 3, 2.75,
                  0.24, 3, 3, 3, 3, 1.8]  # 欲設置的欄位高度(cm)
    for i in range(len(row_height)):
        row_height[i] *= 566.929

    for rows in table.rows:
        set_row_height(rows, row_height[rows._index])  # 設置儲存格高度為row_height

    table.cell(0, 0).text = df.columns.names[0]  # 添加索引名稱

    row_name_1 = ['8:10 | 9:00', '9:10 | 10:00', '10:10 | 11:00', '11:10 | 12:00',
                  '12:40 13:30', '1:40 | 2:30', '2:40 | 3:30', '3:40 | 4:30', '4:40 | 5:30', '5:40 | 6:30']
    row_name_2 = ['1', '2', '3', '4', 'DN', '5', '6', '7', '8', 'E0']

    for j in range(12):  # 添加列標題(rows):時間
        if j == 1:
            table.cell(j, 0).text = "時間"
        if j > 1:
            table.cell(j, 0).text = row_name_1[j-2]

    for j in range(12):  # 添加列標題(rows):節次
        if j == 1:
            table.cell(j, 1).text = "節次"
        if j > 1:
            table.cell(j, 1).text = row_name_2[j-2]

    days = ['一', '二', '三', '四', '五']
    grades = ['資一', '應一', '資二', '應二', '資三', '應三', '資四', '應四']
    for k in range(2):  # 添加欄標題(columns):[星期，開課單位]
        for i in range(5):
            for j in range(8):
                if k == 0:
                    m = 8*i+j+2
                    table.cell(k, m).text = days[i]
                if k == 1:
                    m = 8*i+j+2
                    table.cell(k, m).text = grades[j]

    # 合併星期標題
    for k in range(5):
        start = table.cell(0, 8*k+2)
        end = table.cell(0, 8*k+2+7)
        start_text = start.text
        start.merge(end)
        start.text = start_text
        start.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 將必修課填入表格
    for i in range(df.shape[1]):
        if df.iloc[2, i] == 0:
            continue
        else:
            day = df.columns[i][0]
            grade = df.columns[i][1]
            hour = df.iloc[1, i]  # 課程時間
            start = hour[0]
            end = hour[-1]
            header_row_1 = table.rows[0]
            k = -1
            for cell_1 in header_row_1.cells:
                k += 1
                if cell_1.text != day:
                    continue
                for n in range(8):
                    if grade == table.cell(1, k+n).text:
                        a = table.cell(start+2, k+n)
                        b = table.cell(end+2, k+n)
                        a.merge(b)
                        table.cell(start+2, k+n).text = df.iloc[0, i]
                        break
                break

    # 加入剩下的課程
    for i in range(df.shape[1]):
        if df.iloc[2, i] == 1:
            continue
        else:
            day = df.columns[i][0]
            grade = df.columns[i][1]
            hour = df.iloc[1, i]  # 課程時間
            start = hour[0]
            end = hour[-1]
            header_row_1 = table.rows[0]
            k = -1
            for cell_1 in header_row_1.cells:
                if (df.iloc[2, i] == -1 or df.iloc[2, i] == 1):
                    break
                k += 1
                if cell_1.text != day:
                    continue
                for n in range(8):
                    if df.iloc[2, i] == 1 or df.iloc[2, i] == -1:
                        break
                    if grade == table.cell(1, k+n).text:
                        flag = 0
                        for m in range(len(hour)):
                            if len(table.cell(2+hour[m], k+n).text) != 0:
                                flag = 1
                                break
                        if flag == 1:
                            df.iloc[2, i] = -1
                            break

                        if flag == 0:
                            a = table.cell(start+2, k+n)
                            b = table.cell(end+2, k+n)
                            a.merge(b)
                            table.cell(start+2, k+n).text = df.iloc[0, i]
                            df.iloc[2, i] = 1
                            break

    # 無法順利將字體設為標楷體
    # def set_cell_font(cell, font_name='標楷體', font_size=11):  # 表格中的文字設定
    #    paragraph = document.add_paragraph(cell.text)
    #    text = cell.text
    #    run = paragraph.add_run(text)
    #    paragraphs = cell.paragraphs
    #    for paragraph in paragraphs:
    #        for run in paragraph.runs:  # 獲取段落中的所有文字 run
    #            run.font.name = font_name  # 設置文字 run 的字型名稱
    #            if font_size:
    #                run.font.size = Pt(font_size)  # 設置文字 run 的字型大小

    # for row in table.rows:
    #    for cell in row.cells:
    #        set_cell_font(cell)

    # 定義一個函數來設置表格邊框

    def set_table_borders(table):
        tbl = table._tbl  # 獲取邊框
        tblBorders = OxmlElement('w:tblBorders')

        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')  # 設置邊框樣式
            border.set(qn('w:sz'), '4')        # 設置邊框寬度
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')  # 設置邊框顏色
            tblBorders.append(border)

        tblPr = tbl.tblPr
        tblPr.append(tblBorders)

    set_table_borders(table)

    #    for char in text:
    #        p = cell.add_paragraph(char)
    #        run = p.runs[0]
    #        run.font.size = Pt(12)  # 調整字體大小
    #        run.font.name = '標楷體'  # 設置字體為標楷體

    # def set_text_vertical(cell):
    #    tc = cell._tc
    #    tcPr = tc.get_or_add_tcPr()
    #    text_direction = OxmlElement('w:textDirection')
    #    text_direction.set(qn('w:val'), 'tbRl')  # tbRl 是從上到下的直書方向
    #    tcPr.append(text_direction)

    # for row in table.rows:
    #    for cell in row.cells:
    #        set_text_vertical(cell)

    # for i in range(42):  # 添加欄標題(columns)
    #    for j in range(2):
    #        index_value = df.columns[i][j]
    #        table.cell(j, i+2).text = index_value

    # for i in range(df.shape[0]):
    #    for j in range(df.shape[1]):
    #        table.cell(i + 2, j + 1).text = str(df.iloc[i, j])

    document.save(filename)  # 儲存文件


current_directory = os.path.dirname(os.path.abspath(__file__))
course_file_path = os.path.join(current_directory, '數學系課表.docx')
dataframe_to_word(df_math, course_file_path)
# dataframe_to_word(df_math, r"C:\Users\User\OneDrive\Desktop\排課程式\數學系課表.docx")
course_fill_file_path = os.path.join(current_directory, '未填入課程(-1未填入).xlsx')
df_math.to_excel(course_fill_file_path, index=True)
# df_math.to_excel(r"C:\Users\User\OneDrive\Desktop\排課程式\數學系課表測試3.xlsx", index=True)

# print(df_math['一', '資數一'])  # column name也可具有多個層級，如同multiindex

# syllabus=pd.DataFrame(index=time,columns=)


# def course_schedule(df):

# data = {'age': ['2', '3', '5', '7', '11'],
#        'Name': ['Alice', 'Bob', 'Charlie', 'David', 'Emma']}
# df = pd.DataFrame(data)
# df['Name'] = '(' + df['Name'].astype(str) + ')'
